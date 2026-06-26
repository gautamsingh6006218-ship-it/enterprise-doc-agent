"""
ingestion/loaders/docx.py

What problem does this solve?
- Word documents (.docx) are one of the most common enterprise formats.
  Raw .docx files are ZIP archives containing XML — unreadable without a parser.

Why python-docx?
- Most mature open-source library for .docx parsing. Handles all paragraph
  styles (headings, body text, lists) uniformly.

What is NOT extracted (intentional for this phase):
- Tables           → need row/column-aware chunking — Phase 2.
- Headers/footers  → usually page numbers and company names, not content.
                     Including them adds noise to embeddings.
- Embedded images  → OCR pipeline — Phase 2.
- Tracked changes  → python-docx surfaces accepted text only, not diffs.

Note on .doc (legacy Word 97-2003 binary format):
- .doc requires antiword or LibreOffice CLI to convert.
- Not supported — LoaderRegistry returns a clear error for .doc files.
"""

import hashlib
from pathlib import Path
from uuid import uuid4

from docx import Document as DocxDocument

from agent.ingestion.loaders.base import BaseDocumentLoader
from agent.ingestion.models import Document


class DocxLoader(BaseDocumentLoader):
    """
    What problem does this solve?
    - Converts a .docx file into plain text by extracting all body paragraphs.

    Why does this class exist?
    - Isolates python-docx dependency in one place. Swapping the library or
      adding table extraction only touches this file.
    """

    @property
    def supported_extensions(self) -> list[str]:
        """Handles .docx only. Legacy .doc binary format is not supported."""
        return [".docx"]

    def load(self, file_path: str, tenant_id: str = "default") -> Document:
        """
        What problem does this solve?
        - Turns a Word document into indexed, searchable plain text.

        Why are these inputs required?
        - file_path:  Location of the .docx file on disk.
        - tenant_id:  Propagated to Document for tenant-scoped vector queries.

        Why filter empty paragraphs?
        - Word uses blank paragraphs for visual spacing (line breaks between
          sections). Including them injects whitespace-only strings into chunks,
          wasting embedding capacity and degrading retrieval quality.

        Why join paragraphs with double newline?
        - Preserves paragraph boundaries as natural split points for the chunker.
          Single \n would merge separate paragraphs into one long block.

        Why paragraph_count in metadata?
        - Gives the chunker a rough sense of document size before splitting.
          Also useful for debugging when a document produces fewer chunks than expected.

        Why Document instead of str?
        - Same as all loaders: downstream services need source path, tenant,
          hash, and metadata alongside the text.

        Raises:
        - FileNotFoundError  if the file does not exist.
        - ValueError         if the file is not a .docx.
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        if path.suffix.lower() not in self.supported_extensions:
            raise ValueError(f"DocxLoader cannot handle: {path.suffix}")

        raw_bytes = path.read_bytes()
        file_hash = hashlib.md5(raw_bytes).hexdigest()

        doc = DocxDocument(str(path))

        # Blank paragraphs are spacing artefacts in Word, not content.
        # Filtering them keeps the text corpus clean for embedding.
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n\n".join(paragraphs).strip()

        return Document(
            id=str(uuid4()),
            source_type="docx",
            source_path=str(path.resolve()),
            title=path.stem,
            text=full_text,
            tenant_id=tenant_id,
            file_hash=file_hash,
            metadata={
                "paragraph_count": len(paragraphs),
                "file_name": path.name,
                "file_size_bytes": len(raw_bytes),
            },
        )
